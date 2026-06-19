from __future__ import annotations

import csv
import json
from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parent
OUT = ROOT / "FDP_EVC_Project_Handoff.docx"

BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
INK = RGBColor(15, 23, 42)
MUTED = RGBColor(85, 85, 85)
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F2F4F7"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_width(cell, width_dxa: int) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths_dxa: list[int], indent_dxa: int = 120) -> None:
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths_dxa)))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent_dxa))
    tbl_ind.set(qn("w:type"), "dxa")

    grid = tbl.tblGrid
    if grid is None:
        grid = OxmlElement("w:tblGrid")
        tbl.insert(0, grid)
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(width))
        grid.append(grid_col)

    for row in table.rows:
        for i, cell in enumerate(row.cells):
            set_cell_width(cell, widths_dxa[i])
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_cell_margins(table, top=80, start=120, bottom=80, end=120) -> None:
    tbl_pr = table._tbl.tblPr
    tbl_cell_mar = tbl_pr.find(qn("w:tblCellMar"))
    if tbl_cell_mar is None:
        tbl_cell_mar = OxmlElement("w:tblCellMar")
        tbl_pr.append(tbl_cell_mar)
    for side, value in [("top", top), ("start", start), ("bottom", bottom), ("end", end)]:
        node = tbl_cell_mar.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            tbl_cell_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_run_font(run, size=None, color=None, bold=None, italic=None, name="Calibri") -> None:
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def add_para(doc, text="", *, style=None, bold=False, italic=False, color=None, size=None, align=None):
    p = doc.add_paragraph(style=style)
    if align is not None:
        p.alignment = align
    run = p.add_run(text)
    set_run_font(run, size=size, color=color, bold=bold, italic=italic)
    return p


def add_bullets(doc, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.left_indent = Inches(0.375)
        p.paragraph_format.first_line_indent = Inches(-0.188)
        p.paragraph_format.space_after = Pt(4)
        run = p.add_run(item)
        set_run_font(run)


def add_numbers(doc, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p.paragraph_format.left_indent = Inches(0.375)
        p.paragraph_format.first_line_indent = Inches(-0.188)
        p.paragraph_format.space_after = Pt(4)
        run = p.add_run(item)
        set_run_font(run)


def add_heading(doc, text: str, level: int = 1):
    p = doc.add_paragraph(style=f"Heading {level}")
    run = p.add_run(text)
    if level == 1:
        set_run_font(run, size=16, color=BLUE, bold=True)
    elif level == 2:
        set_run_font(run, size=13, color=BLUE, bold=True)
    else:
        set_run_font(run, size=12, color=DARK_BLUE, bold=True)
    return p


def add_note(doc, label: str, text: str, fill: str = "F4F6F9") -> None:
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [9360], indent_dxa=120)
    set_cell_margins(table, top=120, bottom=120, start=160, end=160)
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    p = cell.paragraphs[0]
    r = p.add_run(f"{label}: ")
    set_run_font(r, bold=True, color=DARK_BLUE)
    r = p.add_run(text)
    set_run_font(r)
    doc.add_paragraph()


def add_table(doc, headers: list[str], rows: list[list[str]], widths: list[int], font_size: float = 9):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    set_table_geometry(table, widths)
    set_cell_margins(table)
    for i, header in enumerate(headers):
        cell = table.cell(0, i)
        set_cell_shading(cell, LIGHT_BLUE)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(header)
        set_run_font(r, size=font_size, bold=True, color=INK)
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            p = cells[i].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT if i == 0 else WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(str(value))
            set_run_font(r, size=font_size)
            cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    doc.add_paragraph()
    return table


def add_image(doc, path: str, caption: str, width: float = 6.2) -> None:
    img = ROOT / path
    if not img.exists():
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(img), width=Inches(width))
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = cap.add_run(caption)
    set_run_font(r, size=9, color=MUTED, italic=True)


def fmt(x: float) -> str:
    return f"{float(x):.3f}"


def read_coat_csv(label: str) -> list[list[str]]:
    path = ROOT / "dog" / "experiments" / "coat" / label / "test_metrics.csv"
    out = []
    with path.open(newline="", encoding="utf-8") as fh:
        reader = csv.DictReader(fh)
        for row in reader:
            out.append([
                row["method"],
                fmt(row["pr_auc"]),
                fmt(row["roc_auc"]),
                fmt(row["f1"]),
                fmt(row["precision"]),
                fmt(row["recall"]),
                fmt(row["threshold"]),
            ])
    return out


def read_eye_rows() -> list[list[str]]:
    exp = ROOT / "dog" / "experiments" / "eye"
    base = json.loads((exp / "baseline_results.json").read_text(encoding="utf-8"))
    rows = []
    for method in ["Majority", "LR", "RF"]:
        t = base[method]["test"]
        rows.append([method, fmt(t["pr_auc"]), fmt(t["roc_auc"]), fmt(t["f1"]), fmt(t["precision"]), fmt(t["recall"])])
    files = [
        ("MLP", "mlp_results.json"),
        ("MLP tuned", "mlp_best_results.json"),
        ("TabPFN", "tabpfn_results.json"),
        ("TabICL", "tabicl_results.json"),
        ("TabNet", "tabnet_results.json"),
    ]
    for method, filename in files:
        t = json.loads((exp / filename).read_text(encoding="utf-8"))["test"]
        rows.append([method, fmt(t["pr_auc"]), fmt(t["roc_auc"]), fmt(t["f1"]), fmt(t["precision"]), fmt(t["recall"])])
    return rows


def setup_document() -> Document:
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for level, size, color, before, after in [
        (1, 16, BLUE, 18, 10),
        (2, 13, BLUE, 14, 7),
        (3, 12, DARK_BLUE, 10, 5),
    ]:
        style = styles[f"Heading {level}"]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.25

    header = section.header.paragraphs[0]
    header.text = "FDP-EVC Project Handoff"
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_run_font(header.runs[0], size=9, color=MUTED)
    footer = section.footer.paragraphs[0]
    footer.text = "Internal handoff guide - generated from project source"
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run_font(footer.runs[0], size=8, color=MUTED)
    return doc


def build() -> None:
    doc = setup_document()

    # First-page memo masthead.
    add_para(doc, "FDP-EVC PROJECT HANDOFF", bold=True, color=INK, size=23)
    add_para(
        doc,
        "Tài liệu bàn giao từ đầu đến cuối cho người chưa đọc project",
        color=MUTED,
        size=13,
    )
    add_table(
        doc,
        ["Mục", "Thông tin"],
        [
            ["Project", "FDP-EVC - Externally Visible Characteristics from DNA"],
            ["Phạm vi", "Part A: Human HIrisPlex silver labels; Part B: Dog eye/coat ground-truth labels"],
            ["Ngày chốt", date.today().isoformat()],
            ["Mục tiêu người đọc", "Hiểu project đủ để viết báo cáo, trình bày kết quả, và biết các caveat cần nói rõ"],
        ],
        [1800, 7560],
        font_size=9,
    )
    add_note(
        doc,
        "Thông điệp chính",
        "Project đi từ bài toán human với nhãn silver do HIrisPlex sinh ra sang dog với nhãn phenotype thật. MLP là model đề xuất thực dụng: chạy được full protocol, có kết quả cạnh tranh, và dễ giải thích hơn foundation models nặng như TabPFN/TabICL.",
        fill="E8EEF5",
    )

    add_heading(doc, "1. Cách đọc nhanh tài liệu này", 1)
    add_numbers(
        doc,
        [
            "Đọc mục 2 để hiểu thesis chung của project.",
            "Đọc mục 3 nếu cần viết phần human: dataset, silver labels, EDA, baselines, MLP M1/M2.",
            "Đọc mục 4 nếu cần viết phần dog: eye, coat, GWAS feature selection, model comparison.",
            "Đọc mục 5 để biết cách giải thích vì sao diagram cho ra kết quả như vậy.",
            "Đọc mục 6-7 để biết cách chạy lại và những điều không nên claim trong báo cáo.",
        ],
    )

    add_heading(doc, "2. Narrative tổng thể của project", 1)
    add_para(
        doc,
        "FDP-EVC nghiên cứu dự đoán đặc điểm nhìn thấy được từ DNA. Project được chia thành hai phần vì hai loại nhãn khác nhau tạo ra hai câu chuyện khác nhau.",
    )
    add_table(
        doc,
        ["Phần", "Dữ liệu", "Loại nhãn", "Câu hỏi nghiên cứu", "Kết luận chính"],
        [
            [
                "Human",
                "2,481 người, 41 HIrisPlex SNP",
                "Silver labels từ HIrisPlex-S",
                "Điều gì xảy ra khi train lại model trên nhãn do model khác sinh ra?",
                "Accuracy cao nhưng macro-F1 bị giới hạn; argmax làm mất class hiếm. Soft-label KL giúp skin tốt hơn.",
            ],
            [
                "Dog",
                "Eye: 2,769 dogs; Coat: 1,930 dogs",
                "Ground truth tương đối trực tiếp: ảnh/survey",
                "Khi có phenotype thật, MLP và tabular foundation models hoạt động ra sao?",
                "MLP cạnh tranh và chạy full protocol; TabPFN/TabICL mạnh nhưng rất nặng với coat.",
            ],
        ],
        [900, 1700, 1600, 2500, 2660],
        font_size=8.5,
    )

    add_heading(doc, "3. Part A - Human HIrisPlex pipeline", 1)
    add_heading(doc, "3.1 Dataset và preprocessing", 2)
    add_para(
        doc,
        "Nguồn human gồm `hirisplex_results_FN_v2.csv` và `full_dataset.csv`. File HIrisPlex chứa mỗi sample với chuỗi `input_csv`: dòng 1 là tên SNP, dòng 2 là dosage 0/1/2 của 41 SNP. Sau khi bỏ 23 sample thiếu input, pipeline tạo dataset ML với 2,481 samples x 41 SNP.",
    )
    add_bullets(
        doc,
        [
            "Features: `snp_0` đến `snp_40`, mỗi SNP là dosage 0/1/2.",
            "Labels: eye có 3 lớp, hair có 6 lớp, skin có 5 lớp.",
            "Nhãn hard ban đầu lấy bằng argmax trên xác suất/p_value HIrisPlex trong từng nhóm trait.",
            "Persistent split: 80% trainval, 20% test; trong trainval có 5-fold CV với seed 42.",
        ],
    )
    add_heading(doc, "3.2 EDA: vấn đề của argmax silver labels", 2)
    add_table(
        doc,
        ["Trait", "Phân bố argmax label", "Vấn đề"],
        [
            ["Eye", "blue 58.0%, intermediate 0%, brown 42.0%", "Class intermediate biến mất hoàn toàn."],
            ["Hair", "brown 6.0%, red 0.2%, light 32.9%, dark 60.9%", "Blond/black gần như không xuất hiện; red cực hiếm."],
            ["Skin", "intermediate 44.9%, dark 6.3%, dark_to_black 48.4%", "Very pale/pale gần như mất; ordinal structure bị bỏ qua."],
        ],
        [1200, 3100, 5060],
        font_size=8.5,
    )
    add_table(
        doc,
        ["Threshold", "Eye low-conf", "Hair low-conf", "Skin low-conf"],
        [
            ["0.5", "0.44%", "0.00%", "5.24%"],
            ["0.6", "61.19%", "5.84%", "12.45%"],
            ["0.7", "72.23%", "22.49%", "21.77%"],
            ["0.8", "76.30%", "60.98%", "33.90%"],
        ],
        [1600, 2500, 2500, 2760],
        font_size=9,
    )
    add_image(doc, "human/report/figures/01_label_distribution.png", "Human Figure 1 - phân bố nhãn argmax.", 6.2)

    add_heading(doc, "3.3 Baselines và MLP", 2)
    add_para(
        doc,
        "Baselines gồm Logistic Regression, SVM-RBF, và Random Forest. MLP có hai biến thể: M1 dùng CrossEntropy trên hard argmax label; M2 dùng KL-divergence trên soft label HIrisPlex probability vector.",
    )
    add_table(
        doc,
        ["Model", "Loss/ý tưởng", "Điểm cần nhớ"],
        [
            ["LR/SVM/RF", "Train trên hard labels", "Là baseline sạch nhưng chỉ học lại output của HIrisPlex."],
            ["M1 MLP CE", "CrossEntropy với label argmax", "Cô lập hiệu ứng architecture MLP so với baseline."],
            ["M2 MLP KL", "KL(q_HIrisPlex || p_model)", "Không ép sample low-confidence thành một nhãn cứng; tận dụng phân phối xác suất."],
        ],
        [1600, 3000, 4760],
        font_size=9,
    )
    add_para(doc, "Công thức quan trọng:")
    add_bullets(
        doc,
        [
            "Hard CE: L = -log p_model(y_argmax | x).",
            "Soft KL: L = sum_c q_c log(q_c / p_c), trong đó q là phân phối HIrisPlex và p là phân phối model.",
            "Với ordinal traits như hair/skin, MAE và QWK có ý nghĩa hơn accuracy vì khoảng cách giữa lớp được quan tâm.",
        ],
    )
    add_table(
        doc,
        ["Method", "Eye macro-F1", "Hair macro-F1", "Skin macro-F1", "Điểm giải thích"],
        [
            ["Best baseline", "~0.667", "0.545", "0.652", "Accuracy cao nhưng macro-F1 thấp do class collapse."],
            ["M1 MLP CE", "0.665", "0.552", "0.615", "MLP không tự sửa được lỗi nhãn hard argmax."],
            ["M2 MLP KL", "0.667", "0.545", "0.774", "Soft label giúp skin rõ nhất vì giữ thông tin phân phối thay vì chỉ argmax."],
        ],
        [1600, 1500, 1500, 1500, 3260],
        font_size=8.5,
    )
    add_image(doc, "human/report/figures/08_method_comparison.png", "Human Figure 2 - so sánh test macro-F1 giữa baselines và MLP.", 6.2)

    add_heading(doc, "3.4 Kết luận phần human", 2)
    add_bullets(
        doc,
        [
            "Human part không chứng minh model dự đoán phenotype thật; nó chứng minh pipeline ML và giới hạn của silver labels.",
            "Nhãn HIrisPlex là model output, nên mọi model train từ đó bị giới hạn bởi chính HIrisPlex.",
            "Soft-label KL là improvement hợp lý vì giữ uncertainty, đặc biệt khi argmax label low-confidence.",
            "Đây là lý do chuyển sang dog: cần phenotype quan sát trực tiếp để vượt trần silver labels.",
        ],
    )

    doc.add_page_break()
    add_heading(doc, "4. Part B - Dog ground-truth phenotype pipeline", 1)
    add_heading(doc, "4.1 Vì sao chuyển sang dog", 2)
    add_para(
        doc,
        "Dog part xử lý cùng loại bài toán genotype -> externally visible characteristic, nhưng nhãn không phải output của một model khác. Eye labels đến từ owner/photo verification; coat labels đến từ Darwin's Ark survey. Vì vậy dog là phần dùng để kiểm tra model trên phenotype thật hơn.",
    )

    add_heading(doc, "4.2 Dog eye color", 2)
    add_bullets(
        doc,
        [
            "Dataset: Deane-Coe et al. 2018, n = 2,769 dogs.",
            "Task: binary classification, blue eyes = 1, brown/non-blue = 0.",
            "Positive rate chỉ khoảng 3.9%, nên accuracy không dùng làm metric chính.",
            "Raw genotype ban đầu có 213,245 SNP; pipeline chọn 52 SNP với `p_wald < 5e-8` theo ngưỡng genome-wide significance trong paper.",
            "Phần lớn SNP nằm quanh chr18/ALX4, đúng locus sinh học liên quan blue eyes.",
        ],
    )
    add_note(
        doc,
        "SNP dosage 0/1/2",
        "Mỗi SNP là một vị trí biến thể. Chó có hai bản sao chromosome nên genotype được mã hóa thành số bản sao của allele được đếm: 0, 1, hoặc 2. Ma trận model là rows = dogs, columns = SNPs.",
        fill="F4F6F9",
    )
    add_heading(doc, "4.3 Models trong dog eye", 2)
    add_table(
        doc,
        ["Model", "Cách học", "Vai trò trong báo cáo"],
        [
            ["Majority", "Luôn đoán class phổ biến", "Mốc sàn; PR-AUC bằng positive rate."],
            ["LR/RF", "Classical ML", "Baseline dễ giải thích."],
            ["MLP", "2 hidden layers, BCEWithLogitsLoss + pos_weight", "Model đề xuất thực dụng."],
            ["MLP tuned", "Grid search, chọn bằng CV PR-AUC", "Kiểm tra default MLP có hợp lý không."],
            ["TabPFN", "Pretrained Transformer, in-context learning", "Foundation baseline mạnh nhưng nặng."],
            ["TabICL", "Column-wise embedding, row-wise interaction, dataset-wise ICL", "Foundation model mới, hợp tabular context."],
            ["TabNet", "Sequential feature attention, train từ đầu", "So sánh attention kiểu chọn feature."],
        ],
        [1300, 3800, 4260],
        font_size=8.2,
    )
    add_table(
        doc,
        ["Model", "PR-AUC", "ROC-AUC", "F1", "Precision", "Recall"],
        read_eye_rows(),
        [1800, 1200, 1200, 1000, 1200, 1200],
        font_size=8.5,
    )
    add_image(doc, "dog/report/figures/02_eye_metric_bars.png", "Dog Figure 1 - eye-color test metrics trên 52 SNP.", 6.2)
    add_heading(doc, "4.4 Giải thích kết quả eye", 3)
    add_bullets(
        doc,
        [
            "MLP tuned và TabICL có PR-AUC gần nhau nhất: 0.670 vs 0.673. Nghĩa là khả năng ranking positive cao hơn negative gần tương đương.",
            "TabICL có F1 cao nhất vì threshold validation của nó là 0.08, phù hợp rare positive; F1 rất nhạy vì test chỉ có khoảng 22 positive.",
            "TabNet F1 thấp không có nghĩa nó không học: ROC-AUC vẫn 0.817, nhưng threshold 0.5 làm precision rất thấp và tạo nhiều false positives.",
            "Kết luận đúng: MLP ổn và cạnh tranh; không nên claim MLP thắng tuyệt đối mọi model.",
        ],
    )

    add_heading(doc, "4.5 Dog coat color", 2)
    add_para(
        doc,
        "Coat dùng Darwin's Ark Q243 survey. Đây là multi-label phenotype: một dog có thể có nhiều màu cùng lúc. Project hiện báo cáo hai binary tasks chính: black coat và red/liver/brown/tan coat.",
    )
    add_table(
        doc,
        ["Màu/nhóm màu", "Số dog", "Tỉ lệ"],
        [
            ["white_or_cream", "1,377", "71.3%"],
            ["white", "1,194", "61.9%"],
            ["black", "1,057", "54.8%"],
            ["red/liver/brown/tan", "937", "48.5%"],
            ["tan", "561", "29.1%"],
            ["cream", "417", "21.6%"],
            ["liver/brown", "365", "18.9%"],
            ["yellow", "319", "16.5%"],
            ["red", "232", "12.0%"],
            ["grey/blue", "205", "10.6%"],
        ],
        [3200, 1800, 1800],
        font_size=9,
    )
    add_para(
        doc,
        "Feature selection cho coat đọc từng GWAS output file trong Dryad zip. Mỗi binary label có một GWAS file riêng. Với `p < 5e-8`, black có 719 SNP, red_brown_tan có 526 SNP.",
    )
    add_table(
        doc,
        ["Task", "Samples", "Positive rate", "SNPs", "Protocol"],
        [
            ["black", "1,930", "54.8%", "719", "Classical/MLP/TabNet full CV+test; TabPFN/TabICL test-only do RAM."],
            ["red_brown_tan", "1,930", "48.5%", "526", "Classical/MLP/TabNet full CV+test; TabPFN/TabICL test-only do RAM."],
        ],
        [1600, 1200, 1500, 1000, 4060],
        font_size=8.5,
    )

    add_heading(doc, "4.6 Black coat results", 2)
    add_table(
        doc,
        ["Model", "PR-AUC", "ROC-AUC", "F1", "Precision", "Recall", "Threshold"],
        read_coat_csv("black"),
        [1650, 1050, 1050, 900, 1150, 1150, 1100],
        font_size=8.2,
    )
    add_image(doc, "dog/report/figures/coat_black_metric_bars.png", "Dog Figure 2 - black coat test metrics.", 6.2)

    add_heading(doc, "4.7 Red/brown/tan coat results", 2)
    add_table(
        doc,
        ["Model", "PR-AUC", "ROC-AUC", "F1", "Precision", "Recall", "Threshold"],
        read_coat_csv("red_brown_tan"),
        [1650, 1050, 1050, 900, 1150, 1150, 1100],
        font_size=8.2,
    )
    add_image(doc, "dog/report/figures/coat_red_brown_tan_metric_bars.png", "Dog Figure 3 - red/brown/tan coat test metrics.", 6.2)

    add_heading(doc, "5. Cơ sở lý thuyết để giải thích diagram", 1)
    add_heading(doc, "5.1 Các metric được tính như thế nào", 2)
    add_bullets(
        doc,
        [
            "PR-AUC đo ranking theo precision-recall; phù hợp khi positive hiếm, ví dụ eye chỉ 3.9% blue.",
            "ROC-AUC đo khả năng phân biệt positive/negative trên mọi threshold; ít nhạy hơn PR-AUC khi imbalance mạnh.",
            "F1 = 2TP / (2TP + FP + FN), phụ thuộc threshold hard prediction nên có thể dao động mạnh dù PR-AUC gần như không đổi.",
            "Precision = TP/(TP+FP), recall = TP/(TP+FN). Với rare class, vài mẫu sai có thể làm F1 thay đổi lớn.",
        ],
    )
    add_heading(doc, "5.2 Vì sao MLP là model đề xuất ổn", 2)
    add_bullets(
        doc,
        [
            "Feature selection bằng GWAS đã loại phần lớn nhiễu, nên input còn lại chứa tín hiệu sinh học mạnh hơn.",
            "MLP học nonlinear combinations giữa SNPs, ví dụ SNP_A chỉ quan trọng khi SNP_B cũng xuất hiện.",
            "Class-weighted BCE (`pos_weight = n_neg/n_pos`) giúp rare positive không bị model bỏ qua.",
            "MLP chạy được full CV/test trên cả eye và coat, nên protocol sạch hơn TabPFN/TabICL coat.",
            "Kết quả không luôn tốt nhất tuyệt đối, nhưng cạnh tranh và dễ bảo vệ về mặt thực nghiệm.",
        ],
    )
    add_heading(doc, "5.3 Vì sao TabPFN/TabICL có thể cao nhưng rất nặng", 2)
    add_bullets(
        doc,
        [
            "Hai model này là pretrained tabular foundation models. Chúng không train weights bằng gradient trên dog dataset như MLP.",
            "`fit()` chủ yếu đặt train samples + labels vào context; khi predict, model dùng attention để test sample học từ context.",
            "CV không phải để train lại weights; CV là đánh giá ổn định. Mỗi fold vẫn phải chạy context mới nên rất tốn RAM/CPU.",
            "Black coat có 719 SNP và red_brown_tan có 526 SNP; TabPFN còn cảnh báo khi feature > 500. Vì vậy coat foundation results được chạy test-only.",
        ],
    )
    add_heading(doc, "5.4 Vì sao TabNet không luôn tốt", 2)
    add_bullets(
        doc,
        [
            "TabNet train từ đầu, không pretrained, nên cần đủ dữ liệu positive để học ổn định.",
            "Attention của TabNet là feature selection trong từng sample, khác với TabPFN/TabICL attention theo context giữa samples.",
            "Trên eye, threshold 0.5 làm TabNet recall cao nhưng precision thấp, F1 rơi mạnh.",
            "Trên coat, TabNet hợp hơn eye vì positive rate cân bằng hơn, nhưng vẫn không vượt RF/MLP ổn định.",
        ],
    )
    add_heading(doc, "5.5 Vì sao Majority black có F1 cao", 2)
    add_para(
        doc,
        "Black positive rate là 54.8%. Majority đoán tất cả là black nên recall = 1.0, precision = 0.546, F1 = 0.707. Đây không phải model tốt; ROC-AUC = 0.5 chứng minh nó không phân biệt được gì.",
    )

    doc.add_page_break()
    add_heading(doc, "6. Cách chạy lại project", 1)
    add_heading(doc, "6.1 Human", 2)
    add_table(
        doc,
        ["Bước", "Command", "Output chính"],
        [
            ["Preprocess", "python human/src/data/preprocess.py", "human/data/processed/evc_processed.npz"],
            ["Splits", "python human/src/data/splits.py", "trainval/test/fold npy"],
            ["EDA", "python human/src/data/eda.py", "human/report/eda.md + figures"],
            ["Baselines", "python human/src/models/baselines.py", "baseline_results.json"],
            ["MLP", "python human/src/train/train_mlp.py", "m1_results.json, m2_results.json"],
            ["Reports", "python human/src/evaluation/report_*.py", "baselines.md, mlp.md, figures"],
        ],
        [1400, 4300, 3660],
        font_size=8.2,
    )
    add_heading(doc, "6.2 Dog eye", 2)
    add_table(
        doc,
        ["Bước", "Command", "Output chính"],
        [
            ["Preprocess", "python dog/src/data/preprocess_eye.py", "eye_processed.npz với 52 SNP"],
            ["Splits", "python dog/src/data/splits.py", "eye_splits.json"],
            ["Baselines", "python dog/src/models/baselines.py", "baseline_results.json"],
            ["MLP", "python dog/src/train/train_eye.py", "mlp_results.json"],
            ["MLP tuned", "python dog/src/experiments/tune_mlp.py", "mlp_best_results.json"],
            ["TabPFN/TabICL/TabNet", "python dog/src/train/train_*.py", "tabpfn/tabicl/tabnet results"],
            ["Report", "python dog/src/evaluation/report_eye.py", "eye.md + figures"],
        ],
        [1400, 4300, 3660],
        font_size=8.2,
    )
    add_heading(doc, "6.3 Dog coat", 2)
    add_table(
        doc,
        ["Bước", "Command", "Ghi chú"],
        [
            ["Preprocess black", "python dog/src/data/preprocess_coat.py --label black", "Tạo 719 SNP với p < 5e-8."],
            ["Preprocess red/brown/tan", "python dog/src/data/preprocess_coat.py --label red_brown_tan", "Tạo 526 SNP với p < 5e-8."],
            ["Splits", "python dog/src/data/splits_coat.py --label <label>", "80/20 + 5-fold stratified."],
            ["Full models", "python dog/src/train/train_coat.py --label <label> --models Majority LR RF MLP TabNet", "Chạy full CV/test."],
            ["MLP tuned", "python dog/src/experiments/tune_mlp_coat.py --label <label>", "24 configs x 5 folds."],
            ["Foundation test-only", "python dog/src/experiments/run_coat_foundation_test_only.py --label <label> --model TabPFN", "Dùng khi full CV làm tràn RAM."],
            ["Report", "python dog/src/evaluation/report_coat.py --label <label>", "CSV + metric bar PNG."],
        ],
        [1400, 4300, 3660],
        font_size=8.1,
    )

    add_heading(doc, "7. Những điều phải nói rõ trong báo cáo", 1)
    add_table(
        doc,
        ["Claim", "Có nên nói?", "Cách nói đúng"],
        [
            ["Human model dự đoán phenotype thật.", "Không", "Human dùng silver labels từ HIrisPlex, nên chỉ đánh giá khả năng học lại/khai thác output HIrisPlex."],
            ["MLP là model tốt nhất tuyệt đối.", "Không", "MLP là model đề xuất thực dụng, cạnh tranh và chạy được full protocol."],
            ["TabPFN/TabICL coat hơn MLP chắc chắn.", "Không", "Coat foundation results là test-only/light setting do giới hạn RAM; dùng làm tham khảo mạnh."],
            ["Accuracy là metric chính cho eye.", "Không", "Eye imbalanced 3.9% positive; dùng PR-AUC chính, F1/ROC-AUC phụ."],
            ["p < 1.15e-7 là threshold paper eye.", "Không", "Đã sửa về `p_wald < 5e-8`, ngưỡng genome-wide significance trong paper."],
        ],
        [2500, 1300, 5560],
        font_size=8.2,
    )

    add_heading(doc, "8. Checklist cho người viết báo cáo", 1)
    add_bullets(
        doc,
        [
            "Mở đầu bằng vấn đề EVC từ genotype và sự khác nhau giữa silver labels vs ground-truth labels.",
            "Human: nhấn mạnh EDA cho thấy argmax collapse classes; M2 soft-label KL là cải tiến hợp lý.",
            "Dog eye: giải thích GWAS p<5e-8, imbalance, PR-AUC, MLP với pos_weight.",
            "Dog coat: giải thích multi-label Q243, chọn black và red/brown/tan vì positive rate gần cân bằng và có đủ GWAS signal.",
            "Khi so model, tách rõ full CV/test vs test-only foundation models.",
            "Phần kết luận: MLP là lựa chọn thực dụng, không phải model thắng tuyệt đối; foundation models mạnh nhưng tốn tài nguyên.",
        ],
    )

    add_heading(doc, "9. File quan trọng trong repo", 1)
    add_table(
        doc,
        ["Đường dẫn", "Ý nghĩa"],
        [
            ["human/report/eda.md", "EDA human: phân bố label, low-confidence, vấn đề argmax."],
            ["human/report/baselines.md", "Kết quả LR/SVM/RF human."],
            ["human/report/mlp.md", "Kết quả M1/M2 MLP human và giải thích soft-label KL."],
            ["dog/README.md", "Tóm tắt dog eye pipeline và kết quả chính."],
            ["dog/report/eye.md", "Bảng kết quả dog eye chi tiết."],
            ["dog/experiments/coat/*/test_metrics.csv", "Bảng metric coat dùng cho diagram."],
            ["dog/report/tabicl_knowledge.md", "Tài liệu lý thuyết TabICL/TabPFN/TabNet đã trao đổi."],
            ["dog/report/ml_course_knowledge.md", "Tóm tắt kiến thức ML từ các lesson và câu hỏi đã hỏi."],
        ],
        [3800, 5560],
        font_size=8.5,
    )

    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build()
